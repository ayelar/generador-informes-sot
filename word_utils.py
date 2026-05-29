# ======================================
# SECCIÓN 3 - Word: verificación cruzada y reemplazo (HÍBRIDA)
# - Body + tablas: método convencional (sin runs)
# - Headers/footers: método robusto (runs)
# - Insensible a tildes y mayúsculas
# - Postproceso: "- ± -" --> "-"
# ======================================
from docx import Document
import re
import unicodedata

def generar_word(word_path, diccionario):
    # -------------------------------
    # CONFIG
    # -------------------------------
    WORD_TEMPLATE_PATH = word_path
    WORD_OUTPUT_PATH = "Informe_generado.docx"
    
    doc = Document(WORD_TEMPLATE_PATH)

    # --------------------------------------------------
    # Normalización (tildes, mayúsculas, signos)
    # --------------------------------------------------
    def normalizar(texto):
        if texto is None:
            return ""
        texto = str(texto).strip().lower()

        texto = unicodedata.normalize("NFD", texto)
        texto = "".join(c for c in texto if unicodedata.category(c) != "Mn")

        texto = texto.rstrip(":").rstrip(".").strip()
        return texto

    # --------------------------------------------------
    # Postproceso global de texto (corrige - ± -)
    # --------------------------------------------------
    def postprocesar_texto(texto):
        if texto is None:
            return ""

        texto = str(texto)

        # Convertir NBSP (espacio no separable) a espacio normal
        texto = texto.replace("\u00A0", " ")

        # Reemplazar cualquier variante de "- ± -" por "-"
        # (acepta espacios, tabs, etc.)
        texto = re.sub(r"-\s*±\s*-", "-", texto)

        return texto

    # --------------------------------------------------
    # Diccionario normalizado (Excel sin {})
    # --------------------------------------------------
    diccionario_normalizado = {}
    for clave, valor in diccionario.items():
        diccionario_normalizado[normalizar(clave)] = valor

    # --------------------------------------------------
    # PATRÓN: {clave}
    # --------------------------------------------------
    PATRON_CLAVE = re.compile(r"\{([^}]+)\}")

    # --------------------------------------------------
    # Reemplazo (método convencional) -> SOLO body/tablas
    # --------------------------------------------------
    def reemplazar_convencional_en_parrafo(parrafo, diccionario_norm):
        texto_original = parrafo.text
        if not texto_original:
            return

        texto = texto_original

        def reemplazo(match):
            contenido = match.group(1).strip()
            contenido_limpio = contenido.rstrip(":").rstrip(".").strip()
            clave_norm = normalizar(contenido_limpio)

            if clave_norm in diccionario_norm:
                valor = diccionario_norm[clave_norm]
                if valor == "__BORRAR__":
                    return ""
                return str(valor)

            return match.group(0)

        texto = PATRON_CLAVE.sub(reemplazo, texto)

        # Postproceso "- ± -" -> "-"
        texto = postprocesar_texto(texto)

        if texto != texto_original:
            parrafo.clear()
            parrafo.add_run(texto)

    # --------------------------------------------------
    # Reemplazo robusto en párrafo (runs) -> SOLO headers/footers
    # --------------------------------------------------
    def reemplazar_runs_en_parrafo(parrafo, diccionario_norm):
        texto_original = "".join(run.text for run in parrafo.runs)
        if not texto_original:
            return

        texto = texto_original

        def reemplazo(match):
            contenido = match.group(1).strip()
            contenido_limpio = contenido.rstrip(":").rstrip(".").strip()
            clave_norm = normalizar(contenido_limpio)

            if clave_norm in diccionario_norm:
                valor = diccionario_norm[clave_norm]
                if valor == "__BORRAR__":
                    return ""
                return str(valor)

            return match.group(0)

        texto = PATRON_CLAVE.sub(reemplazo, texto)

        # Postproceso "- ± -" -> "-"
        texto = postprocesar_texto(texto)

        if texto == texto_original:
            return

        # Vaciar runs
        for run in parrafo.runs:
            run.text = ""

        # Escribir en el primer run
        if parrafo.runs:
            parrafo.runs[0].text = texto
        else:
            parrafo.add_run(texto)

    # --------------------------------------------------
    # Procesar tablas
    # --------------------------------------------------
    def procesar_tabla(tabla, diccionario_norm, modo="convencional"):
        for fila in tabla.rows:
            for celda in fila.cells:
                for p in celda.paragraphs:
                    if modo == "convencional":
                        reemplazar_convencional_en_parrafo(p, diccionario_norm)
                    else:
                        reemplazar_runs_en_parrafo(p, diccionario_norm)

                # Tablas anidadas
                for t in celda.tables:
                    procesar_tabla(t, diccionario_norm, modo=modo)

    # --------------------------------------------------
    # Procesar bloque (body/header/footer)
    # --------------------------------------------------
    def procesar_bloque(bloque, diccionario_norm, modo="convencional"):
        for p in bloque.paragraphs:
            if modo == "convencional":
                reemplazar_convencional_en_parrafo(p, diccionario_norm)
            else:
                reemplazar_runs_en_parrafo(p, diccionario_norm)

        for tabla in bloque.tables:
            procesar_tabla(tabla, diccionario_norm, modo=modo)

    # --------------------------------------------------
    # Extraer claves {…} para auditoría (usa p.text)
    # --------------------------------------------------
    def extraer_claves_word(documento):
        claves = set()

        def extraer_de_bloque(bloque):
            for p in bloque.paragraphs:
                claves.update(re.findall(r"\{[^}]+\}", p.text))

            for tabla in bloque.tables:
                for fila in tabla.rows:
                    for celda in fila.cells:
                        for p in celda.paragraphs:
                            claves.update(re.findall(r"\{[^}]+\}", p.text))

        # Body
        extraer_de_bloque(documento)

        # Headers/footers
        for section in documento.sections:
            extraer_de_bloque(section.header)
            extraer_de_bloque(section.footer)

            extraer_de_bloque(section.first_page_header)
            extraer_de_bloque(section.first_page_footer)

            extraer_de_bloque(section.even_page_header)
            extraer_de_bloque(section.even_page_footer)

        return claves

    # --------------------------------------------------
    # VERIFICACIÓN CRUZADA Word ↔ Excel
    # --------------------------------------------------
    claves_word_raw = extraer_claves_word(doc)

    claves_word_norm = set(normalizar(c.strip("{}")) for c in claves_word_raw)
    claves_excel_norm = set(diccionario_normalizado.keys())

    faltantes_en_excel = sorted(claves_word_norm - claves_excel_norm)
    no_usadas_en_word = sorted(claves_excel_norm - claves_word_norm)

    print("\n================ VERIFICACIÓN CRUZADA ================\n")

    print(">>> Claves PRESENTES en la plantilla Word y AUSENTES en el diccionario Excel:")
    if faltantes_en_excel:
        for c in faltantes_en_excel:
            print(f"  - {c}")
    else:
        print("  (Ninguna)")

    print("\n>>> Claves DEFINIDAS en el diccionario Excel y NO ENCONTRADAS en la plantilla Word:")
    if no_usadas_en_word:
        for c in no_usadas_en_word:
            print(f"  - {c}")
    else:
        print("  (Ninguna)")

    print("\n=====================================================\n")

    # --------------------------------------------------
    # REEMPLAZO
    # --------------------------------------------------

    # 1) BODY + TABLAS: convencional
    procesar_bloque(doc, diccionario_normalizado, modo= "runs")

    # 2) HEADERS/FOOTERS: robusto (runs)
    for section in doc.sections:
        procesar_bloque(section.header, diccionario_normalizado, modo="runs")
        procesar_bloque(section.footer, diccionario_normalizado, modo="runs")

        procesar_bloque(section.first_page_header, diccionario_normalizado, modo="runs")
        procesar_bloque(section.first_page_footer, diccionario_normalizado, modo="runs")

        procesar_bloque(section.even_page_header, diccionario_normalizado, modo="runs")
        procesar_bloque(section.even_page_footer, diccionario_normalizado, modo="runs")

    for p in list(doc.paragraphs):

        texto = p.text.strip()


        # borrar líneas con marcador
        if "__BORRAR__" in texto:
            p._element.getparent().remove(p._element)
            continue

        # borrar líneas tipo ":"
        if texto == ":":
            p._element.getparent().remove(p._element)
            continue

        # borrar líneas tipo "M2:"
        if texto.startswith("M") and texto.endswith(":"):
            p._element.getparent().remove(p._element)
            continue

    # --------------------------------------------------
    # LIMPIEZA DE TABLAS
    # Borra filas donde TODOS los resultados son "-"
    # --------------------------------------------------

    for tabla in doc.tables:

        filas_a_borrar = []

        # recorrer filas
        for i, fila in enumerate(tabla.rows):

            # Saltar encabezado
            if i == 0:
                continue

            # Tomar SOLO columnas de resultados
            celdas_resultados = fila.cells[1:]

            resultados = []

            for celda in celdas_resultados:

                texto = celda.text.strip()

                resultados.append(texto)

            # Si TODOS son "-"
            if all(r == "-" or r == "" for r in resultados):

                filas_a_borrar.append(fila)

        # borrar filas
        for fila in filas_a_borrar:

            tabla._tbl.remove(fila._tr)

    # --------------------------------------------------
    # ELIMINAR TABLAS VACÍAS
    # (solo quedó encabezado)
    # --------------------------------------------------

    tablas_a_borrar = []

    for tabla in doc.tables:

        # si solo tiene 1 fila -> encabezado
        if len(tabla.rows) <= 1:
            tablas_a_borrar.append(tabla)

    # eliminar tablas
    for tabla in tablas_a_borrar:

        tabla._element.getparent().remove(tabla._element)

    # Guardar
    doc.save(WORD_OUTPUT_PATH)

    print("Documento Word generado correctamente:")
    print(f"  -> {WORD_OUTPUT_PATH}")

    return WORD_OUTPUT_PATH